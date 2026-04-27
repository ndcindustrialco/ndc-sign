# -*- coding: utf-8 -*-
"""Generate NDC-Sign User Manual (Thai) as a .docx file.

Run: python scripts/build_user_manual.py
Output: public/manual ndc-sign/NDC-Sign_User_Manual.docx
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ------------------------------------------------------------------
# Paths
# ------------------------------------------------------------------
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(ROOT, "public", "manual ndc-sign")
OUT_PATH = os.path.join(IMG_DIR, "NDC-Sign_User_Manual.docx")

FONT_NAME = "Sarabun"
NAVY = RGBColor(0x0B, 0x2E, 0x6F)  # หัวข้อสีน้ำเงินเข้ม
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x4B, 0x55, 0x63)


# ------------------------------------------------------------------
# Low-level helpers
# ------------------------------------------------------------------
def _set_run_font(run, size: float, bold: bool = False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    # ฟอนต์ไทยใน Word ต้องเซ็ต cs (complex script) ด้วย
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    # complex-script size
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    # แทนที่ ถ้ามีแล้ว
    existing = rpr.find(qn("w:szCs"))
    if existing is not None:
        rpr.remove(existing)
    rpr.append(sz_cs)
    run.bold = bold
    if bold:
        bcs = OxmlElement("w:bCs")
        existing_b = rpr.find(qn("w:bCs"))
        if existing_b is not None:
            rpr.remove(existing_b)
        rpr.append(bcs)
    if color is not None:
        run.font.color.rgb = color


def _add_bottom_border(paragraph, color_hex: str = "0B2E6F", size: int = 12):
    """เพิ่มเส้นขีดล่างสีน้ำเงินเข้มใต้หัวข้อ"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    # ลบ pBdr เดิม ถ้ามี
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def _set_paragraph_spacing(paragraph, before: float = 6, after: float = 6, line: float = 1.4):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc: Document, text: str, level: int = 1):
    """หัวข้อมีเส้นขีดล่างสีน้ำเงินเข้ม"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    size = {1: 20, 2: 16, 3: 14}.get(level, 14)
    _set_run_font(run, size=size, bold=True, color=NAVY)
    _set_paragraph_spacing(p, before=14, after=6, line=1.3)
    _add_bottom_border(p, "0B2E6F", size=12 if level == 1 else 8)
    # Bookmark-style: ใช้ outline level เพื่อให้ TOC เก็บได้
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(outline)
    # ใช้ style heading เพื่อให้ TOC (field) จับได้
    p.style = doc.styles[f"Heading {level}"]
    return p


def add_sub_label(doc: Document, text: str):
    """ฉลากเล็ก เช่น 'วัตถุประสงค์' 'ขั้นตอนการใช้งาน'"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=14, bold=True, color=NAVY)
    _set_paragraph_spacing(p, before=8, after=2, line=1.3)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=14, color=TEXT)
    _set_paragraph_spacing(p, before=2, after=2, line=1.5)
    return p


def add_step(doc: Document, index: int, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{index})  {text}")
    _set_run_font(run, size=14, color=TEXT)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(-0.8)
    _set_paragraph_spacing(p, before=1, after=1, line=1.5)
    return p


def add_note(doc: Document, items: list[str]):
    for it in items:
        p = doc.add_paragraph()
        run = p.add_run(f"•  {it}")
        _set_run_font(run, size=14, color=TEXT)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.8)
        pf.first_line_indent = Cm(-0.8)
        _set_paragraph_spacing(p, before=1, after=1, line=1.5)


def add_image(doc: Document, filename: str, caption: str, width_cm: float = 15.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        run = p.add_run(f"[ไม่พบภาพ: {filename}]")
        _set_run_font(run, size=12, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    _set_paragraph_spacing(p, before=6, after=2, line=1.0)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    _set_run_font(run, size=12, color=MUTED)
    run.italic = True
    _set_paragraph_spacing(cap, before=0, after=10, line=1.2)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------------------
# Document setup
# ------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()

    # ตั้งค่าฟอนต์ default
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # Heading styles
    for lvl, size in [(1, 20), (2, 16), (3, 14)]:
        hstyle = doc.styles[f"Heading {lvl}"]
        hstyle.font.name = FONT_NAME
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        hstyle.font.color.rgb = NAVY
        hrpr = hstyle.element.get_or_add_rPr()
        hrFonts = hrpr.find(qn("w:rFonts"))
        if hrFonts is None:
            hrFonts = OxmlElement("w:rFonts")
            hrpr.append(hrFonts)
        hrFonts.set(qn("w:ascii"), FONT_NAME)
        hrFonts.set(qn("w:hAnsi"), FONT_NAME)
        hrFonts.set(qn("w:cs"), FONT_NAME)
        hrFonts.set(qn("w:eastAsia"), FONT_NAME)

    # กระดาษ A4 และระยะขอบ
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    return doc


# ------------------------------------------------------------------
# Cover page
# ------------------------------------------------------------------
def build_cover(doc: Document):
    # เว้นพื้นที่ด้านบนเล็กน้อย
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("คู่มือการใช้งานระบบ")
    _set_run_font(r, size=28, bold=True, color=NAVY)
    _set_paragraph_spacing(title, before=0, after=4, line=1.2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("NDC-Sign")
    _set_run_font(r, size=36, bold=True, color=NAVY)
    _set_paragraph_spacing(sub, before=0, after=4, line=1.1)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("ระบบลงนามเอกสารอิเล็กทรอนิกส์")
    _set_run_font(r, size=18, color=TEXT)
    _set_paragraph_spacing(sub2, before=0, after=40, line=1.2)

    # เส้นคั่น
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(line_p, "0B2E6F", size=18)
    _set_paragraph_spacing(line_p, before=0, after=30, line=1.0)

    for _ in range(6):
        doc.add_paragraph()

    meta1 = doc.add_paragraph()
    meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta1.add_run("หน่วยงานจัดทำ")
    _set_run_font(r, size=14, color=MUTED)
    _set_paragraph_spacing(meta1, before=0, after=2, line=1.2)

    meta1v = doc.add_paragraph()
    meta1v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta1v.add_run("ฝ่ายเทคโนโลยีสารสนเทศ บริษัท เอ็น ดี ซี อีเลคทรอนิคส์ จำกัด")
    _set_run_font(r, size=16, bold=True, color=TEXT)
    _set_paragraph_spacing(meta1v, before=0, after=20, line=1.2)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2.add_run("วันที่ปรับปรุงล่าสุด")
    _set_run_font(r, size=14, color=MUTED)
    _set_paragraph_spacing(meta2, before=0, after=2, line=1.2)

    meta2v = doc.add_paragraph()
    meta2v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2v.add_run("23 เมษายน 2569")
    _set_run_font(r, size=16, bold=True, color=TEXT)
    _set_paragraph_spacing(meta2v, before=0, after=2, line=1.2)

    meta2e = doc.add_paragraph()
    meta2e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2e.add_run("(Version 1.0)")
    _set_run_font(r, size=12, color=MUTED)
    _set_paragraph_spacing(meta2e, before=0, after=0, line=1.2)

    add_page_break(doc)


# ------------------------------------------------------------------
# คำนำ
# ------------------------------------------------------------------
def build_preface(doc: Document):
    add_heading(doc, "คำนำ", level=1)
    add_body(
        doc,
        "เอกสารฉบับนี้จัดทำขึ้นเพื่อเป็นแนวทางการใช้งานระบบ NDC-Sign ซึ่งเป็นระบบลงนามเอกสาร"
        "อิเล็กทรอนิกส์ที่ฝ่ายเทคโนโลยีสารสนเทศพัฒนาขึ้นสำหรับใช้ภายในองค์กร โดยรองรับการอัปโหลด"
        "เอกสาร การวางช่องลงนาม การกำหนดผู้ลงนามและลำดับการลงนาม การติดตามสถานะ "
        "ตลอดจนการจัดเก็บหลักฐานการลงนามที่ตรวจสอบย้อนหลังได้",
    )
    add_body(
        doc,
        "ผู้ใช้งานควรศึกษาขั้นตอนตามลำดับบท เพื่อให้เข้าใจการทำงานของระบบอย่างครบถ้วน "
        "หากพบปัญหาระหว่างใช้งาน กรุณาส่งข้อเสนอแนะผ่านเมนูความคิดเห็นในระบบ "
        "เพื่อให้ผู้พัฒนาปรับปรุงระบบให้ใช้งานได้ดียิ่งขึ้น",
    )
    add_body(doc, "ฝ่ายเทคโนโลยีสารสนเทศ")
    add_body(doc, "บริษัท เอ็น ดี ซี อีเลคทรอนิคส์ จำกัด")
    add_page_break(doc)


# ------------------------------------------------------------------
# สารบัญ (แบบ static)
# ------------------------------------------------------------------
def build_toc(doc: Document):
    add_heading(doc, "สารบัญ", level=1)

    entries = [
        ("บทที่ 1  การเข้าสู่ระบบ", None),
        ("        1.1  การลงชื่อเข้าใช้ด้วยบัญชีองค์กร", None),
        ("บทที่ 2  หน้าหลักและภาพรวมเอกสาร", None),
        ("        2.1  หน้าแดชบอร์ด", None),
        ("บทที่ 3  การอัปโหลดเอกสาร", None),
        ("        3.1  การนำเข้าไฟล์ PDF", None),
        ("บทที่ 4  การวางช่องลงนามในเอกสาร", None),
        ("        4.1  การเพิ่มและปรับแต่งช่อง", None),
        ("บทที่ 5  การกำหนดผู้ลงนาม", None),
        ("        5.1  การเพิ่มผู้ลงนามด้วยอีเมล", None),
        ("        5.2  ลำดับการลงนามกรณีหลายคน", None),
        ("        5.3  หัวข้อและข้อความแจ้งผู้ลงนาม", None),
        ("บทที่ 6  การติดตามสถานะเอกสาร", None),
        ("        6.1  สถานะเปิดดู", None),
        ("        6.2  สถานะลงนามแล้ว", None),
        ("        6.3  รายละเอียดเอกสาร", None),
        ("        6.4  การส่งแจ้งเตือนซ้ำ", None),
        ("บทที่ 7  การใช้งานฝั่งผู้ลงนาม", None),
        ("        7.1  หน้าลงนามของผู้อนุมัติ", None),
        ("        7.2  การบันทึกลายเซ็นเพื่อใช้ครั้งถัดไป", None),
        ("        7.3  การปฏิเสธการลงนามพร้อมเหตุผล", None),
        ("บทที่ 8  เมื่อการลงนามเสร็จสมบูรณ์", None),
        ("        8.1  อีเมลแจ้งผลและไฟล์ตรวจสอบ", None),
        ("บทที่ 9  ช่องทางข้อเสนอแนะ", None),
        ("        9.1  การส่งความคิดเห็นและแจ้งปัญหา", None),
    ]
    for text, _ in entries:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_run_font(run, size=14, color=TEXT)
        _set_paragraph_spacing(p, before=2, after=2, line=1.3)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.0)
    add_page_break(doc)


# ------------------------------------------------------------------
# Section block (หัวข้อแต่ละเรื่อง: วัตถุประสงค์ / ขั้นตอน / ข้อพึงทราบ / ภาพ)
# ------------------------------------------------------------------
def add_topic(
    doc: Document,
    *,
    heading: str,
    purpose: str,
    steps: list[str],
    notes: list[str],
    image_file: str,
    caption: str,
    level: int = 2,
    image_width: float = 15.0,
):
    add_heading(doc, heading, level=level)

    add_sub_label(doc, "วัตถุประสงค์")
    add_body(doc, purpose)

    add_sub_label(doc, "ขั้นตอนการใช้งาน")
    for i, s in enumerate(steps, start=1):
        add_step(doc, i, s)

    add_sub_label(doc, "ข้อพึงทราบ")
    add_note(doc, notes)

    add_sub_label(doc, "ภาพประกอบ")
    add_image(doc, image_file, caption, width_cm=image_width)


# ------------------------------------------------------------------
# Chapters
# ------------------------------------------------------------------
def build_chapters(doc: Document):
    # ---------- บทที่ 1 ----------
    add_heading(doc, "บทที่ 1  การเข้าสู่ระบบ", level=1)
    add_topic(
        doc,
        heading="1.1  การลงชื่อเข้าใช้ด้วยบัญชีองค์กร",
        purpose="เพื่อให้พนักงานเข้าสู่ระบบ NDC-Sign ด้วยบัญชีไมโครซอฟท์ขององค์กร และยืนยันตัวตนก่อนใช้งาน",
        steps=[
            "เปิดหน้าเว็บ NDC-Sign ผ่านเบราว์เซอร์",
            "กดปุ่ม “ลงชื่อเข้าใช้ด้วย Microsoft”",
            "กรอกอีเมลและรหัสผ่านของบริษัท แล้วยืนยันตัวตนตามขั้นตอน",
            "ระบบนำเข้าสู่หน้าแดชบอร์ดโดยอัตโนมัติเมื่อยืนยันสำเร็จ",
        ],
        notes=[
            "ใช้ได้เฉพาะบัญชีพนักงานที่องค์กรอนุญาตเท่านั้น",
            "หากเข้าใช้ไม่ได้ ให้ตรวจสอบการเชื่อมต่ออินเทอร์เน็ตและสิทธิ์การเข้าถึง",
            "ระบบจะออกจากระบบให้อัตโนมัติเมื่อไม่มีการใช้งานเป็นเวลานาน",
        ],
        image_file="sign-in.png",
        caption="ภาพที่ 1-1  หน้าลงชื่อเข้าใช้ระบบ NDC-Sign",
    )
    add_page_break(doc)

    # ---------- บทที่ 2 ----------
    add_heading(doc, "บทที่ 2  หน้าหลักและภาพรวมเอกสาร", level=1)
    add_topic(
        doc,
        heading="2.1  หน้าแดชบอร์ด",
        purpose="เพื่อดูเอกสารทั้งหมดของผู้ใช้ในที่เดียว พร้อมสถานะและเมนูการทำงานหลัก",
        steps=[
            "เลือกเมนู “เอกสาร” จากแถบด้านซ้าย",
            "ดูรายการเอกสารที่แสดงในตารางพร้อมคอลัมน์สถานะ",
            "กรองรายการตามสถานะ เช่น ฉบับร่าง รอลงนาม เสร็จสิ้น",
            "ค้นหาเอกสารด้วยชื่อเรื่องได้จากช่องค้นหา",
            "กดที่ชื่อเอกสารเพื่อเปิดดูรายละเอียด",
        ],
        notes=[
            "ผู้ใช้แต่ละคนจะเห็นเฉพาะเอกสารของตนเองหรือเอกสารที่เกี่ยวข้อง",
            "สถานะจะเปลี่ยนอัตโนมัติเมื่อมีการลงนามหรือเปิดดูเอกสาร",
        ],
        image_file="Dashboard.png",
        caption="ภาพที่ 2-1  หน้าแดชบอร์ดแสดงรายการเอกสาร",
    )
    add_page_break(doc)

    # ---------- บทที่ 3 ----------
    add_heading(doc, "บทที่ 3  การอัปโหลดเอกสาร", level=1)
    add_topic(
        doc,
        heading="3.1  การนำเข้าไฟล์ PDF",
        purpose="เพื่อเพิ่มเอกสารเข้าสู่ระบบก่อนเริ่มกำหนดช่องลงนามและส่งให้ผู้ลงนาม",
        steps=[
            "กดปุ่ม “อัปโหลดเอกสาร” ที่หน้าแดชบอร์ด",
            "ลากไฟล์ PDF มาวางในกรอบ หรือกดเลือกไฟล์จากเครื่อง",
            "ตั้งชื่อเอกสารให้ชัดเจน",
            "กดปุ่มยืนยันเพื่อเข้าสู่หน้าวางช่องลงนาม",
        ],
        notes=[
            "รองรับเฉพาะไฟล์นามสกุล PDF ขนาดไม่เกิน 10 เมกะไบต์",
            "ควรตรวจสอบเนื้อหาไฟล์ให้ครบถ้วนก่อนอัปโหลด",
            "ชื่อเอกสารจะปรากฏในอีเมลถึงผู้ลงนาม",
        ],
        image_file="upload.png",
        caption="ภาพที่ 3-1  หน้าอัปโหลดเอกสาร",
    )
    add_page_break(doc)

    # ---------- บทที่ 4 ----------
    add_heading(doc, "บทที่ 4  การวางช่องลงนามในเอกสาร", level=1)
    add_topic(
        doc,
        heading="4.1  การเพิ่มและปรับแต่งช่อง",
        purpose="เพื่อกำหนดตำแหน่งและประเภทของช่องให้ผู้ลงนามกรอกข้อมูล เช่น ลายเซ็น วันที่ ข้อความ",
        steps=[
            "เลือกประเภทช่องจากแถบด้านข้าง เช่น ลายเซ็น วันที่ ข้อความ",
            "คลิกตำแหน่งบนเอกสารเพื่อวางช่องลงบนหน้าที่ต้องการ",
            "กำหนดผู้รับผิดชอบช่องด้วยการเลือกกลุ่มสี",
            "ตั้งค่าเป็นช่องบังคับหรือไม่บังคับตามความเหมาะสม",
            "ปรับขนาดและตำแหน่งโดยลากกรอบช่อง",
        ],
        notes=[
            "ช่องบังคับต้องกรอกก่อน จึงจะส่งลงนามได้",
            "แต่ละสีแทนผู้ลงนามคนละคน เพื่อแยกหน้าที่ให้ชัดเจน",
            "หากต้องการลบช่อง ให้เลือกช่องแล้วกดไอคอนถังขยะ",
        ],
        image_file="edit field.png",
        caption="ภาพที่ 4-1  หน้าแก้ไขและเพิ่มช่องลงนาม",
    )
    add_body(doc, "")
    add_image(
        doc,
        "วาง field.png",
        "ภาพที่ 4-2  ตัวอย่างการวางช่องลงนามในเอกสาร",
        width_cm=15.0,
    )
    add_page_break(doc)

    # ---------- บทที่ 5 ----------
    add_heading(doc, "บทที่ 5  การกำหนดผู้ลงนาม", level=1)

    add_topic(
        doc,
        heading="5.1  การเพิ่มผู้ลงนามด้วยอีเมล",
        purpose="เพื่อระบุผู้ที่ต้องลงนามในเอกสาร และให้ระบบค้นหาชื่อที่เกี่ยวข้องในองค์กรให้โดยอัตโนมัติ",
        steps=[
            "กดปุ่ม “ส่งลงนาม” ที่มุมบนขวา",
            "พิมพ์ชื่อหรืออีเมลของผู้ลงนามในช่องค้นหา",
            "เลือกรายชื่อที่ระบบแนะนำ",
            "ตรวจสอบอีเมลให้ถูกต้องก่อนยืนยัน",
        ],
        notes=[
            "ระบบแสดงเฉพาะอีเมลที่อยู่ในบริษัทเท่านั้น",
            "กรณีเพิ่มผู้ลงนามผิดคน ให้กดไอคอนลบแล้วเลือกใหม่",
        ],
        image_file="ส่งลงนามโดยพิมชื่อลงไปจะแสดงอีเมลที่เกี่ยวข้องในบริษัท.png",
        caption="ภาพที่ 5-1  การเพิ่มผู้ลงนามและแสดงอีเมลที่เกี่ยวข้อง",
    )

    add_topic(
        doc,
        heading="5.2  ลำดับการลงนามกรณีหลายคน",
        purpose="เพื่อควบคุมว่าให้ผู้ลงนามทุกคนลงนามพร้อมกัน หรือให้ลงนามตามลำดับเป็นขั้น ๆ",
        steps=[
            "เลือกโหมด “ตามลำดับ” หรือ “พร้อมกัน” ในหน้ากำหนดผู้ลงนาม",
            "จัดลำดับผู้ลงนามโดยลากแถวขึ้นลงตามที่ต้องการ",
            "ตรวจสอบหมายเลขลำดับหน้าชื่อผู้ลงนาม",
            "กดยืนยันเพื่อบันทึกการตั้งค่า",
        ],
        notes=[
            "โหมด “ตามลำดับ” ระบบจะส่งอีเมลให้คนถัดไปเมื่อคนก่อนหน้าลงนามเสร็จ",
            "โหมด “พร้อมกัน” ระบบจะส่งอีเมลให้ทุกคนในเวลาเดียวกัน",
        ],
        image_file="ลำดับการเซ็นในกรณีที่มีคนเซ็นหลายคน.png",
        caption="ภาพที่ 5-2  การตั้งค่าลำดับการลงนามเมื่อมีหลายคน",
    )

    add_topic(
        doc,
        heading="5.3  หัวข้อและข้อความแจ้งผู้ลงนาม",
        purpose="เพื่อปรับข้อความในอีเมลที่จะส่งถึงผู้ลงนาม ให้สื่อสารได้ชัดเจนและเหมาะกับเอกสารแต่ละฉบับ",
        steps=[
            "เปิดขั้นตอน “ข้อความแจ้งเตือน” ในหน้าตั้งค่าการส่ง",
            "แก้ไขหัวข้ออีเมลให้ตรงกับเรื่อง",
            "ปรับเนื้อความแจ้งเตือนตามความเหมาะสม",
            "กดส่งเพื่อแจ้งไปยังผู้ลงนามทั้งหมด",
        ],
        notes=[
            "ระบบมีข้อความตั้งต้นภาษาไทยและอังกฤษให้แก้ไขได้",
            "ข้อความจะปรากฏในอีเมลเชิญลงนามที่ผู้รับเปิดอ่าน",
        ],
        image_file="สามารถปรับแจ่งหัวข้อและข้อความที่จะส่งไปหาผู้เซ็นได้.png",
        caption="ภาพที่ 5-3  การปรับหัวข้อและข้อความอีเมลถึงผู้ลงนาม",
    )
    add_page_break(doc)

    # ---------- บทที่ 6 ----------
    add_heading(doc, "บทที่ 6  การติดตามสถานะเอกสาร", level=1)

    add_topic(
        doc,
        heading="6.1  สถานะเปิดดู",
        purpose="เพื่อให้ผู้ส่งทราบว่าผู้ลงนามได้เปิดอ่านเอกสารแล้ว แม้ยังไม่ได้ลงนาม",
        steps=[
            "เปิดรายละเอียดเอกสารจากหน้าแดชบอร์ด",
            "ดูสถานะผู้ลงนามแต่ละราย",
            "ตรวจสอบเวลาที่ผู้ลงนามเปิดเอกสารครั้งแรก",
        ],
        notes=[
            "สถานะ “เปิดดู” ไม่ได้แปลว่าลงนามสำเร็จ",
            "หากผู้ลงนามเปิดดูแต่ยังไม่ลงนามภายในเวลาที่ต้องการ ให้พิจารณาส่งแจ้งเตือนซ้ำ",
        ],
        image_file="status เปิดดู.png",
        caption="ภาพที่ 6-1  สถานะเอกสารเมื่อผู้ลงนามเปิดดู",
    )

    add_topic(
        doc,
        heading="6.2  สถานะลงนามแล้ว",
        purpose="เพื่อยืนยันว่าผู้ลงนามดำเนินการเสร็จสิ้น และติดตามความคืบหน้าของการลงนามทั้งหมด",
        steps=[
            "เปิดรายละเอียดเอกสาร",
            "ดูสัญลักษณ์ “ลงนามแล้ว” ที่หน้ารายชื่อผู้ลงนาม",
            "ตรวจสอบวันที่และเวลาการลงนาม",
        ],
        notes=[
            "เมื่อทุกคนลงนามครบ ระบบจะเปลี่ยนสถานะเอกสารเป็น “เสร็จสิ้น”",
            "ข้อมูลเวลาลงนามจะถูกบันทึกในบันทึกการตรวจสอบ",
        ],
        image_file="status เซ็นแล้ว.png",
        caption="ภาพที่ 6-2  สถานะเอกสารเมื่อผู้ลงนามลงนามเรียบร้อย",
    )

    add_topic(
        doc,
        heading="6.3  รายละเอียดเอกสาร",
        purpose="เพื่อดูข้อมูลเอกสารแบบละเอียด รายชื่อผู้ลงนาม และลำดับเหตุการณ์ทั้งหมด",
        steps=[
            "กดชื่อเอกสารในหน้าแดชบอร์ด",
            "ดูข้อมูลเอกสาร ผู้สร้าง และวันที่ส่ง",
            "ตรวจสอบรายชื่อผู้ลงนามและสถานะรายบุคคล",
            "ดูลำดับเหตุการณ์ เช่น สร้าง ส่ง เปิดดู ลงนาม",
        ],
        notes=[
            "ลำดับเหตุการณ์จะแสดงชื่อผู้ทำรายการและเวลาเสมอ",
            "ข้อมูลในหน้านี้เป็นหลักฐานอ้างอิงเมื่อมีการตรวจสอบย้อนหลัง",
        ],
        image_file="รายละเอียด Detail.png",
        caption="ภาพที่ 6-3  หน้ารายละเอียดเอกสารและลำดับเหตุการณ์",
    )

    add_topic(
        doc,
        heading="6.4  การส่งแจ้งเตือนซ้ำ",
        purpose="เพื่อกระตุ้นให้ผู้ลงนามที่ยังไม่ดำเนินการ กลับมาลงนามให้เอกสารเสร็จสมบูรณ์",
        steps=[
            "เปิดรายละเอียดเอกสารที่ยังไม่เสร็จ",
            "กดปุ่ม “ส่งแจ้งเตือน” ข้างชื่อผู้ลงนามที่ต้องการ",
            "ยืนยันการส่งในหน้าต่างที่ปรากฏ",
        ],
        notes=[
            "ระบบจะส่งอีเมลฉบับใหม่พร้อมลิงก์ลงนามเดิม",
            "ควรเว้นระยะก่อนส่งซ้ำ เพื่อไม่รบกวนผู้ลงนามมากเกินไป",
        ],
        image_file="ส่งแจ้งเตือน Reminder.png",
        caption="ภาพที่ 6-4  การส่งแจ้งเตือนผู้ลงนามอีกครั้ง",
    )
    add_page_break(doc)

    # ---------- บทที่ 7 ----------
    add_heading(doc, "บทที่ 7  การใช้งานฝั่งผู้ลงนาม", level=1)

    add_topic(
        doc,
        heading="7.1  หน้าลงนามของผู้อนุมัติ",
        purpose="เพื่อให้ผู้ได้รับเอกสารเปิดอ่านและลงนามได้จากลิงก์ในอีเมล โดยไม่ต้องติดตั้งโปรแกรมเพิ่ม",
        steps=[
            "เปิดอีเมลที่ได้รับและกดลิงก์เข้าหน้าเอกสาร",
            "อ่านเนื้อหาเอกสารให้ครบถ้วน",
            "กดช่องลงนามที่ไฮไลต์ไว้",
            "วาดลายเซ็นหรือเลือกลายเซ็นที่บันทึกไว้",
            "กดปุ่ม “ยืนยันการลงนาม”",
        ],
        notes=[
            "รองรับการใช้งานทั้งคอมพิวเตอร์และโทรศัพท์มือถือ",
            "ลิงก์ลงนามใช้ได้ครั้งเดียวและมีวันหมดอายุ",
        ],
        image_file="หน้าผู้อนุมัติ.jpeg",
        caption="ภาพที่ 7-1  หน้าลงนามของผู้อนุมัติ",
    )

    add_topic(
        doc,
        heading="7.2  การบันทึกลายเซ็นเพื่อใช้ครั้งถัดไป",
        purpose="เพื่ออำนวยความสะดวกผู้ลงนามที่ต้องลงนามบ่อย ให้ไม่ต้องวาดลายเซ็นใหม่ทุกครั้ง",
        steps=[
            "วาดลายเซ็นในช่องลงนามครั้งแรก",
            "ติ๊กตัวเลือก “บันทึกลายเซ็นเพื่อใช้ครั้งถัดไป”",
            "กดยืนยันการลงนาม",
            "ครั้งถัดไป ระบบจะเสนอลายเซ็นที่บันทึกไว้ให้เลือกใช้",
        ],
        notes=[
            "ลายเซ็นจะผูกกับบัญชีผู้ลงนาม และเลือกลบได้ในภายหลัง",
            "หากต้องการเปลี่ยน ให้วาดใหม่แล้วบันทึกทับ",
        ],
        image_file="ผู้อนุมัติสามารถบันทึกลายเซ็นตัวเองใช้ครั้งหน้าได้.jpeg",
        caption="ภาพที่ 7-2  การบันทึกลายเซ็นไว้ใช้ในครั้งถัดไป",
    )

    add_topic(
        doc,
        heading="7.3  การปฏิเสธการลงนามพร้อมเหตุผล",
        purpose="เพื่อให้ผู้ลงนามแจ้งเหตุผลเมื่อไม่สามารถลงนามได้ และให้ผู้ส่งนำไปแก้ไขเอกสาร",
        steps=[
            "เปิดเอกสารจากลิงก์ในอีเมล",
            "กดปุ่ม “ปฏิเสธการลงนาม”",
            "พิมพ์เหตุผลการปฏิเสธให้ชัดเจน",
            "กดยืนยันเพื่อส่งกลับถึงผู้ส่ง",
        ],
        notes=[
            "เมื่อปฏิเสธแล้ว จะไม่สามารถลงนามเอกสารฉบับเดิมได้อีก",
            "ผู้ส่งจะได้รับอีเมลแจ้งเหตุผลและต้องส่งเอกสารฉบับใหม่",
        ],
        image_file="ผู้อนุมัติปฏิเสธพร้อมคำอธิบาย.jpeg",
        caption="ภาพที่ 7-3  หน้าจอปฏิเสธการลงนามพร้อมระบุเหตุผล",
    )
    add_page_break(doc)

    # ---------- บทที่ 8 ----------
    add_heading(doc, "บทที่ 8  เมื่อการลงนามเสร็จสมบูรณ์", level=1)
    add_topic(
        doc,
        heading="8.1  อีเมลแจ้งผลและไฟล์ตรวจสอบ",
        purpose="เพื่อส่งมอบเอกสารที่ลงนามครบถ้วนและไฟล์บันทึกการตรวจสอบ ให้ผู้เกี่ยวข้องเก็บไว้เป็นหลักฐาน",
        steps=[
            "ระบบส่งอีเมลถึงผู้ส่งและผู้ลงนามเมื่อทุกคนลงนามครบ",
            "เปิดอีเมลและดาวน์โหลดไฟล์ที่แนบ",
            "ไฟล์แรกคือเอกสารที่ลงนามเรียบร้อย",
            "ไฟล์ที่สองคือบันทึกการตรวจสอบ (Audit)",
        ],
        notes=[
            "ไฟล์บันทึกการตรวจสอบใช้ยืนยันว่าเอกสารไม่ถูกแก้ไขย้อนหลัง",
            "ควรจัดเก็บไฟล์ทั้งสองไว้ในระบบเอกสารขององค์กร",
        ],
        image_file="อีเมลเมื่อเซ็นสำเร็จ จะส่งไฟล์ที่เซ็นแล้วกับไฟล์ Audit มาให้.png",
        caption="ภาพที่ 8-1  อีเมลแจ้งผลการลงนามพร้อมไฟล์เอกสารและไฟล์ตรวจสอบ",
    )
    add_page_break(doc)

    # ---------- บทที่ 9 ----------
    add_heading(doc, "บทที่ 9  ช่องทางข้อเสนอแนะ", level=1)
    add_topic(
        doc,
        heading="9.1  การส่งความคิดเห็นและแจ้งปัญหา",
        purpose="เพื่อให้ผู้ใช้งานแจ้งปัญหาหรือเสนอฟังก์ชันใหม่ และให้ทีมพัฒนานำไปปรับปรุงระบบ",
        steps=[
            "เลือกเมนู “ข้อเสนอแนะ” จากแถบด้านซ้าย",
            "เลือกประเภท เช่น แจ้งปัญหา หรือเสนอฟังก์ชันใหม่",
            "พิมพ์รายละเอียดและแนบภาพประกอบถ้ามี",
            "กดปุ่มส่งเพื่อแจ้งทีมพัฒนา",
        ],
        notes=[
            "ทีมพัฒนาจะพิจารณาทุกข้อเสนอแนะ แต่ไม่สามารถตอบกลับได้ทุกรายการ",
            "ระบุรายละเอียดให้ชัดเจนจะช่วยให้แก้ปัญหาได้รวดเร็ว",
        ],
        image_file="แสดงความคิดเห็น ติดปัญหาหรืออยากให้พเพิ่มอะไร.png",
        caption="ภาพที่ 9-1  หน้าส่งความคิดเห็นและแจ้งปัญหา",
    )


# ------------------------------------------------------------------
# main
# ------------------------------------------------------------------
def main():
    doc = setup_document()
    build_cover(doc)
    build_preface(doc)
    build_toc(doc)
    build_chapters(doc)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
